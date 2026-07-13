import os
import time
import html
import base64
from docx import Document

SECTION_ORDER = [
    "International and Multinational Organisations",
    "G20 and G7",
    "UK and EU Government",
    "Think Tanks",
    "International Health Organisations",
    "Legislative and Specialist Networks",
    "Academic Organisations",
    "Media Organisations",
]

# ---- brand ----
NAVY = "#17457B"
GREEN = "#7DB63F"
KICKER = "#5C8A1F"
LINK = "#1f5a8f"
INK = "#23282b"
MUT = "#6b7075"
FAINT = "#9a9ea2"
HAIR = "#ebe8e2"
PAGE = "#eef1f4"
CARD = "#fdfdfb"
SERIF = "'Newsreader',Georgia,'Times New Roman',serif"
SANS = "-apple-system,'Helvetica Neue',Arial,sans-serif"
LOGO_URL = "https://g20healthpartnership.com/wp-content/uploads/2022/01/G20-G7-Logo-Cropped.jpg"


def _date(it):
    return time.strftime("%d/%m/%y", it["published_parsed"]) if it["published_parsed"] else "n/a"


def _long_date():
    t = time.localtime()
    return time.strftime("%A", t) + " · " + str(t.tm_mday) + time.strftime(" %B %Y", t)


def _logo_data_uri(path="assets/logo.png"):
    try:
        with open(path, "rb") as f:
            return "data:image/png;base64," + base64.b64encode(f.read()).decode()
    except Exception:
        return ""


def _kicker(text):
    return (f'<p style="font-family:{SANS};font-size:11px;letter-spacing:.16em;'
            f'text-transform:uppercase;color:{KICKER};font-weight:500;margin:0">{html.escape(text)}</p>')


def _hair():
    return f'<div style="height:1px;background:{HAIR}"></div>'


def _brief_item(it, first=False):
    title = html.escape(it["title"])
    link = html.escape(it["link"], quote=True)
    synopsis = html.escape(it.get("synopsis", ""))
    source = html.escape(it["source"])
    pad = "padding:0 0 16px" if first else "padding:16px 0"
    row = (f'<div style="{pad}">'
           f'<p style="font-family:{SERIF};font-size:17px;line-height:1.32;margin:0 0 4px">'
           f'<a href="{link}" style="color:{LINK};text-decoration:none">{title}</a></p>'
           f'<p style="font-family:{SANS};font-size:14.5px;line-height:1.6;margin:0;color:{MUT}">'
           f'{synopsis} <span style="color:{FAINT}">{source}</span></p>')
    also = it.get("also", [])
    if also:
        links = ", ".join(
            f'<a href="{html.escape(a["link"], quote=True)}" style="color:{LINK};text-decoration:none">{html.escape(a["source"])}</a>'
            for a in also
        )
        row += f'<p style="font-family:{SANS};font-size:12px;color:{FAINT};margin:6px 0 0">Also covered by {links}</p>'
    return row + "</div>"


def build_email_html(items, note="", lead=None, lead_why=""):
    accepts = [it for it in items if it["relevance"] == "ACCEPT"]
    unsure = [it for it in items if it["relevance"] == "UNSURE"]
    body = [it for it in accepts if it is not lead]

    P = f"padding-left:36px;padding-right:36px"
    out = []
    out.append(f'<!DOCTYPE html><html><head><meta charset="utf-8">'
               f'<meta name="viewport" content="width=device-width, initial-scale=1">'
               f"<style>@import url('https://fonts.googleapis.com/css2?family=Newsreader:ital,opsz,wght@0,6..72,400;0,6..72,500;1,6..72,400&display=swap');</style>"
               f'</head><body style="margin:0;background:{PAGE};font-family:{SANS};color:{INK}">')
    out.append(f'<div style="max-width:600px;margin:24px auto;background:{CARD};border:1px solid #e7e4dd;border-radius:12px;overflow:hidden">')
    out.append(f'<div style="height:3px;background:{GREEN}"></div>')

    # masthead
    logo_html = (f'<img src="{LOGO_URL}" alt="G20 &amp; G7 Health &amp; Development Partnership" '
                 f'style="width:232px;max-width:72%;height:auto">')
    out.append(f'<div style="{P};padding-top:26px;padding-bottom:22px;text-align:center;border-bottom:1px solid {HAIR}">'
               f'{logo_html}'
               f'<div style="width:28px;height:2px;background:{GREEN};margin:18px auto 12px"></div>'
               f'<div style="font-family:{SERIF};font-size:27px;color:{NAVY}">Daily Health Monitoring</div>'
               f'<p style="font-family:{SANS};font-size:11px;letter-spacing:.14em;text-transform:uppercase;color:{FAINT};margin:7px 0 0">{_long_date()}</p>'
               f'</div>')

    # editor's note
    if note:
        out.append(f'<div style="{P};padding-top:24px">'
                   f'<p style="font-family:{SERIF};font-size:16px;line-height:1.72;font-style:italic;color:{MUT};margin:0">{html.escape(note)}</p></div>')

    # big story
    if lead is not None:
        out.append(f'<div style="{P};padding-top:28px">')
        out.append(_kicker("The big story"))
        out.append(f'<h1 style="font-family:{SERIF};font-size:27px;line-height:1.24;font-weight:500;margin:12px 0 12px;color:{NAVY}">'
                   f'<a href="{html.escape(lead["link"], quote=True)}" style="color:{NAVY};text-decoration:none">{html.escape(lead["title"])}</a></h1>')
        out.append(f'<p style="font-family:{SERIF};font-size:16.5px;line-height:1.74;margin:0 0 13px">{html.escape(lead.get("synopsis",""))}</p>')
        if lead_why:
            out.append(f'<p style="font-family:{SERIF};font-size:16.5px;line-height:1.74;margin:0 0 16px">'
                       f'<span style="font-variant:small-caps;letter-spacing:.04em;font-size:14px;color:{FAINT}">Why it matters — </span>{html.escape(lead_why)}</p>')
        out.append(f'<a href="{html.escape(lead["link"], quote=True)}" style="font-family:{SANS};font-size:14px;color:{LINK};text-decoration:none;border-bottom:1px solid #cbe0b3;padding-bottom:2px">Read at {html.escape(lead["source"])}</a>')
        out.append('</div>')

    # in brief, grouped by section
    for section in SECTION_ORDER:
        sec = [it for it in body if it["section"] == section]
        if not sec:
            continue
        out.append(f'<div style="{P};padding-top:28px">{_hair()}</div>')
        out.append(f'<div style="{P};padding-top:20px">{_kicker(section)}</div>')
        out.append(f'<div style="{P};padding-top:14px">')
        for i, it in enumerate(sec):
            if i > 0:
                out.append(_hair())
            out.append(_brief_item(it, first=(i == 0)))
        out.append('</div>')

    # reviewer discretion
    if unsure:
        out.append(f'<div style="{P};padding-top:28px">{_hair()}</div>')
        out.append(f'<div style="{P};padding-top:20px">{_kicker("Potentially relevant — reviewer discretion")}</div>')
        out.append(f'<div style="{P};padding-top:14px">')
        for i, it in enumerate(unsure):
            if i > 0:
                out.append(_hair())
            out.append(_brief_item(it, first=(i == 0)))
        out.append('</div>')

    # footer
    out.append(f'<div style="background:{NAVY};padding:24px 36px;text-align:center;margin-top:28px">'
               f'<div style="height:2px;width:28px;background:{GREEN};margin:0 auto 12px"></div>'
               f'<p style="font-size:11px;letter-spacing:.2em;text-transform:uppercase;color:#a9c6de;margin:0">One message, one voice</p>'
               f'<p style="font-size:12.5px;line-height:1.7;margin:10px 0 0;color:#cfe0ee">'
               f'G20 &amp; G7 Health &amp; Development Partnership · London<br>'
               f'<span style="color:{GREEN}">LinkedIn · X · Web</span></p></div>')

    out.append('</div></body></html>')
    return "\n".join(out)


def build_pdf(html_str, filename):
    """Render the same HTML to a branded PDF. Defensive: if the PDF engine
    isn't installed, skip gracefully so the run still emails."""
    try:
        from weasyprint import HTML
        HTML(string=html_str).write_pdf(filename)
        return filename
    except Exception as ex:
        print(f"  PDF skipped ({type(ex).__name__}: {ex})")
        return None


# ---- legacy Word output (kept so the Streamlit app keeps working) ----
def build_doc(items, filename):
    doc = Document()
    doc.add_heading("Daily Health Monitoring", level=0)
    accepts = [it for it in items if it["relevance"] == "ACCEPT"]
    unsure = [it for it in items if it["relevance"] == "UNSURE"]
    for section in SECTION_ORDER:
        section_items = [it for it in accepts if it["section"] == section]
        if not section_items:
            continue
        doc.add_heading(section, level=1)
        for it in section_items:
            doc.add_heading(f"{it['source']} — {_date(it)} — {it['title']}", level=2)
            doc.add_paragraph(it.get("synopsis", ""))
            doc.add_paragraph(it["link"])
    if unsure:
        doc.add_heading("Potentially relevant — reviewer discretion", level=1)
        for it in unsure:
            doc.add_heading(f"{it['source']} — {_date(it)} — {it['title']}", level=2)
            doc.add_paragraph(it.get("synopsis", ""))
            doc.add_paragraph(it["link"])
    doc.save(filename)
